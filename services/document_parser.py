import os
import io
from typing import Dict, List, Optional, Tuple
import logging
from pathlib import Path
import magic
import tempfile

logger = logging.getLogger(__name__)

class DocumentParser:
    """文档解析服务"""
    
    def __init__(self):
        self.supported_formats = {
            'pdf': self._parse_pdf,
            'docx': self._parse_docx,
            'doc': self._parse_doc,
            'txt': self._parse_txt,
            'html': self._parse_html,
            'xml': self._parse_xml
        }
    
    def parse_document(self, file_data: bytes, filename: str) -> Dict:
        """
        解析文档内容
        
        Args:
            file_data: 文件二进制数据
            filename: 文件名
            
        Returns:
            Dict: 包含解析结果的字典
        """
        try:
            # 检测文件类型
            file_type = self._detect_file_type(file_data, filename)
            logger.info(f"检测到文件类型: {file_type}")
            
            # 检查是否支持该格式
            if file_type not in self.supported_formats:
                logger.warning(f"不支持的文件格式: {file_type}")
                return {
                    'success': False,
                    'error': f'不支持的文件格式: {file_type}',
                    'file_type': file_type,
                    'content': '',
                    'metadata': {}
                }
            
            # 解析文档
            parser_func = self.supported_formats[file_type]
            content, metadata = parser_func(file_data, filename)
            
            return {
                'success': True,
                'file_type': file_type,
                'content': content,
                'metadata': metadata,
                'content_length': len(content),
                'error': None
            }
            
        except Exception as e:
            logger.error(f"解析文档失败: {e}")
            import traceback
            traceback.print_exc()
            return {
                'success': False,
                'error': str(e),
                'file_type': 'unknown',
                'content': '',
                'metadata': {}
            }
    
    def _detect_file_type(self, file_data: bytes, filename: str) -> str:
        """检测文件类型"""
        try:
            # 首先根据文件扩展名判断
            file_ext = Path(filename).suffix.lower().lstrip('.')
            if file_ext in self.supported_formats:
                return file_ext
            
            # 使用magic库检测MIME类型
            mime_type = magic.from_buffer(file_data, mime=True)
            
            # MIME类型到文件类型的映射
            mime_to_type = {
                'application/pdf': 'pdf',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
                'application/msword': 'doc',
                'text/plain': 'txt',
                'text/html': 'html',
                'application/xml': 'xml',
                'text/xml': 'xml'
            }
            
            return mime_to_type.get(mime_type, 'unknown')
            
        except Exception as e:
            logger.error(f"文件类型检测失败: {e}")
            return 'unknown'
    
    def _parse_pdf(self, file_data: bytes, filename: str) -> Tuple[str, Dict]:
        """解析PDF文档"""
        try:
            import PyPDF2
            import pdfplumber
            
            content = ""
            metadata = {"pages": 0, "method": ""}
            
            # 首先尝试使用pdfplumber
            try:
                with pdfplumber.open(io.BytesIO(file_data)) as pdf:
                    metadata["pages"] = len(pdf.pages)
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            content += page_text + "\n"
                    metadata["method"] = "pdfplumber"
                    
            except Exception as e:
                logger.warning(f"pdfplumber解析失败，尝试PyPDF2: {e}")
                
                # 备用方案：使用PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
                metadata["pages"] = len(pdf_reader.pages)
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
                metadata["method"] = "PyPDF2"
            
            return content.strip(), metadata
            
        except ImportError:
            raise ImportError("需要安装PDF解析库: pip install PyPDF2 pdfplumber")
        except Exception as e:
            logger.error(f"PDF解析失败: {e}")
            raise
    
    def _parse_docx(self, file_data: bytes, filename: str) -> Tuple[str, Dict]:
        """解析DOCX文档"""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(file_data))
            content = ""
            metadata = {"paragraphs": 0, "tables": 0}
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
                    metadata["paragraphs"] += 1
            
            # 提取表格文本
            for table in doc.tables:
                metadata["tables"] += 1
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    content += " | ".join(row_text) + "\n"
            
            return content.strip(), metadata
            
        except ImportError:
            raise ImportError("需要安装DOCX解析库: pip install python-docx")
        except Exception as e:
            logger.error(f"DOCX解析失败: {e}")
            raise
    
    def _parse_doc(self, file_data: bytes, filename: str) -> Tuple[str, Dict]:
        """解析DOC文档（旧格式）"""
        try:
            # 使用antiword或其他工具解析DOC文件
            # 这里是一个简化的实现，实际中可能需要更复杂的处理
            import subprocess
            import tempfile
            
            with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                temp_file.write(file_data)
                temp_file.flush()
                
                try:
                    # 尝试使用antiword
                    result = subprocess.run(
                        ['antiword', temp_file.name],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    
                    if result.returncode == 0:
                        content = result.stdout
                        metadata = {"method": "antiword"}
                    else:
                        raise Exception("antiword解析失败")
                        
                except (subprocess.TimeoutExpired, FileNotFoundError, Exception) as e:
                    logger.warning(f"antiword解析失败: {e}, 尝试备用方法")
                    # 备用方案：简单的文本提取
                    content = "DOC文件解析需要专门的工具，当前仅提取基本信息"
                    metadata = {"method": "fallback"}
                
                finally:
                    os.unlink(temp_file.name)
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"DOC解析失败: {e}")
            # 返回基本信息而不是抛出异常
            return f"DOC文件解析失败: {e}", {"method": "error"}
    
    def _parse_txt(self, file_data: bytes, filename: str) -> Tuple[str, Dict]:
        """解析TXT文档"""
        try:
            # 尝试不同的编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
            content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    content = file_data.decode(encoding)
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                # 如果所有编码都失败，使用错误处理模式
                content = file_data.decode('utf-8', errors='ignore')
                used_encoding = 'utf-8-ignore'
            
            metadata = {
                "encoding": used_encoding,
                "lines": len(content.split('\n')),
                "chars": len(content)
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"TXT解析失败: {e}")
            raise
    
    def _parse_html(self, file_data: bytes, filename: str) -> Tuple[str, Dict]:
        """解析HTML文档"""
        try:
            from bs4 import BeautifulSoup
            
            # 尝试解码HTML
            try:
                html_content = file_data.decode('utf-8')
            except UnicodeDecodeError:
                html_content = file_data.decode('gbk', errors='ignore')
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # 移除脚本和样式标签
            for script in soup(["script", "style"]):
                script.decompose()
            
            # 提取文本
            content = soup.get_text()
            # 清理多余的空白字符
            lines = (line.strip() for line in content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            content = ' '.join(chunk for chunk in chunks if chunk)
            
            metadata = {
                "title": soup.title.string if soup.title else "",
                "method": "beautifulsoup"
            }
            
            return content, metadata
            
        except ImportError:
            raise ImportError("需要安装HTML解析库: pip install beautifulsoup4")
        except Exception as e:
            logger.error(f"HTML解析失败: {e}")
            raise
    
    def _parse_xml(self, file_data: bytes, filename: str) -> Tuple[str, Dict]:
        """解析XML文档"""
        try:
            import xml.etree.ElementTree as ET
            
            # 尝试解码XML
            try:
                xml_content = file_data.decode('utf-8')
            except UnicodeDecodeError:
                xml_content = file_data.decode('gbk', errors='ignore')
            
            root = ET.fromstring(xml_content)
            
            # 提取所有文本内容
            content = ""
            for elem in root.iter():
                if elem.text:
                    content += elem.text.strip() + " "
            
            metadata = {
                "root_tag": root.tag,
                "elements_count": len(list(root.iter())),
                "method": "xml.etree"
            }
            
            return content.strip(), metadata
            
        except Exception as e:
            logger.error(f"XML解析失败: {e}")
            raise
    
    def is_suitable_for_knowledge_base(self, content: str, filename: str) -> bool:
        """
        基于简单规则判断文档是否适合加入知识库
        """
        try:
            # 文件名黑名单
            blacklist_keywords = [
                'test', 'temp', 'backup', 'log', 'cache',
                '测试', '临时', '备份', '日志', '缓存'
            ]
            
            filename_lower = filename.lower()
            for keyword in blacklist_keywords:
                if keyword in filename_lower:
                    return False
            
            # 内容长度检查
            if len(content.strip()) < 100:  # 内容太短
                return False
            
            if len(content.strip()) > 100000:  # 内容太长
                return False
            
            # 内容质量检查
            lines = content.split('\n')
            non_empty_lines = [line for line in lines if line.strip()]
            
            if len(non_empty_lines) < 5:  # 有效行数太少
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"适用性检查失败: {e}")
            return False

# 创建全局实例
document_parser = DocumentParser()
